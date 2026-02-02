"""
分节写作模块
根据检索到的内容逐节生成文档
"""
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
from loguru import logger

from .outline_planner import SectionPlan, DocumentOutline


@dataclass
class SectionContent:
    """章节内容"""
    title: str
    level: int
    content: str
    citations: List[Dict[str, Any]]  # 引用列表 [{chunk_id, doc_title, page, quote}]
    word_count: int
    
    def to_dict(self) -> Dict:
        return {
            "title": self.title,
            "level": self.level,
            "content": self.content,
            "citations": self.citations,
            "word_count": self.word_count
        }
    
    def to_markdown(self) -> str:
        """转换为 Markdown"""
        prefix = "#" * (self.level + 1)
        return f"{prefix} {self.title}\n\n{self.content}"


@dataclass
class GeneratedDocument:
    """生成的文档"""
    title: str
    sections: List[SectionContent]
    metadata: Dict[str, Any]
    all_citations: List[Dict[str, Any]]
    
    def to_dict(self) -> Dict:
        return {
            "title": self.title,
            "sections": [s.to_dict() for s in self.sections],
            "metadata": self.metadata,
            "all_citations": self.all_citations
        }
    
    def to_markdown(self) -> str:
        """导出为 Markdown"""
        lines = [f"# {self.title}", ""]
        
        for section in self.sections:
            lines.append(section.to_markdown())
            lines.append("")
        
        # 添加引用列表
        if self.all_citations:
            lines.append("---")
            lines.append("## 参考资料")
            lines.append("")
            
            for i, cite in enumerate(self.all_citations, 1):
                lines.append(f"[{i}] {cite.get('doc_title', '')} - {cite.get('section', '')}, 第{cite.get('page', '')}页")
        
        return "\n".join(lines)
    
    def export(self, path: str, format: str = "markdown"):
        """导出文档"""
        if format == "markdown":
            content = self.to_markdown()
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
        elif format == "docx":
            self._export_docx(path)
        elif format == "html":
            self._export_html(path)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def _export_docx(self, path: str):
        """导出为 Word 文档"""
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            
            # 标题
            doc.add_heading(self.title, 0)
            
            # 章节
            for section in self.sections:
                doc.add_heading(section.title, section.level)
                
                for para in section.content.split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
            
            # 参考资料
            if self.all_citations:
                doc.add_heading("参考资料", 1)
                for i, cite in enumerate(self.all_citations, 1):
                    doc.add_paragraph(
                        f"[{i}] {cite.get('doc_title', '')} - {cite.get('section', '')}"
                    )
            
            doc.save(path)
            logger.info(f"Exported to {path}")
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
    
    def _export_html(self, path: str):
        """导出为 HTML"""
        try:
            import markdown
            
            md_content = self.to_markdown()
            html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
            
            html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{self.title}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
               max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1, h2, h3 {{ color: #333; }}
        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background: #f4f4f4; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
            
            with open(path, "w", encoding="utf-8") as f:
                f.write(html)
            
            logger.info(f"Exported to {path}")
        except ImportError:
            logger.error("markdown not installed. Run: pip install markdown")
            raise


class SectionWriter:
    """
    分节写作器
    
    核心策略:
    1. 每节单独检索所需信息
    2. 基于检索结果生成内容
    3. 强制添加引用标注
    4. 保持术语一致性
    """
    
    def __init__(
        self,
        llm_client=None,
        model: str = "claude-3-5-sonnet-20241022",
        hybrid_searcher=None,
        reranker=None,
        terminology_manager=None
    ):
        self.llm_client = llm_client
        self.model = model
        self.hybrid_searcher = hybrid_searcher
        self.reranker = reranker
        self.terminology_manager = terminology_manager
        
        self._init_llm_client()
    
    def _init_llm_client(self):
        """初始化 LLM 客户端"""
        if self.llm_client is None:
            try:
                # 尝试使用 Anthropic
                from anthropic import Anthropic
                import os
                
                api_key = os.getenv("ANTHROPIC_API_KEY")
                if api_key:
                    self.llm_client = Anthropic(api_key=api_key)
                    self._client_type = "anthropic"
                    return
            except:
                pass
            
            try:
                # 回退到 OpenAI
                from openai import OpenAI
                import os
                
                api_key = os.getenv("OPENAI_API_KEY")
                if api_key:
                    self.llm_client = OpenAI(api_key=api_key)
                    self._client_type = "openai"
                    self.model = "gpt-4.1"
            except Exception as e:
                logger.warning(f"Failed to init LLM client: {e}")
                self._client_type = None
    
    def write_section(
        self,
        section_plan: SectionPlan,
        product: str,
        target_audience: str,
        filter_dict: Optional[Dict[str, Any]] = None,
        top_k: int = 5
    ) -> SectionContent:
        """
        写作单个章节
        
        Args:
            section_plan: 章节规划
            product: 产品名称
            target_audience: 目标读者
            filter_dict: 检索过滤条件
            top_k: 检索数量
        
        Returns:
            SectionContent: 生成的章节内容
        """
        logger.info(f"Writing section: {section_plan.title}")
        
        # 1. 检索相关内容
        contexts, citations = self._retrieve_for_section(
            section_plan=section_plan,
            product=product,
            filter_dict=filter_dict,
            top_k=top_k
        )
        
        # 2. 生成内容
        content = self._generate_section_content(
            section_plan=section_plan,
            contexts=contexts,
            citations=citations,
            product=product,
            target_audience=target_audience
        )
        
        # 3. 术语标准化
        if self.terminology_manager:
            content = self.terminology_manager.normalize_text(content)
        
        word_count = len(content)
        
        return SectionContent(
            title=section_plan.title,
            level=section_plan.level,
            content=content,
            citations=citations,
            word_count=word_count
        )
    
    def write_document(
        self,
        outline: DocumentOutline,
        product: str,
        filter_dict: Optional[Dict[str, Any]] = None
    ) -> GeneratedDocument:
        """
        根据大纲生成完整文档
        
        Args:
            outline: 文档大纲
            product: 产品名称
            filter_dict: 检索过滤条件
        
        Returns:
            GeneratedDocument: 生成的文档
        """
        logger.info(f"Generating document: {outline.title}")
        
        sections = []
        all_citations = []
        citation_index = {}  # 用于去重
        
        for section_plan in outline.sections:
            section = self.write_section(
                section_plan=section_plan,
                product=product,
                target_audience=outline.target_audience,
                filter_dict=filter_dict
            )
            
            sections.append(section)
            
            # 收集引用
            for cite in section.citations:
                cite_key = f"{cite.get('chunk_id', '')}_{cite.get('page', '')}"
                if cite_key not in citation_index:
                    citation_index[cite_key] = len(all_citations) + 1
                    all_citations.append(cite)
        
        return GeneratedDocument(
            title=outline.title,
            sections=sections,
            metadata={
                "product": product,
                "doc_type": outline.doc_type,
                "target_audience": outline.target_audience
            },
            all_citations=all_citations
        )
    
    def _retrieve_for_section(
        self,
        section_plan: SectionPlan,
        product: str,
        filter_dict: Optional[Dict[str, Any]],
        top_k: int
    ) -> tuple:
        """为章节检索相关内容"""
        if not self.hybrid_searcher:
            return [], []
        
        # 构建查询
        queries = [f"{product} {topic}" for topic in section_plan.required_topics]
        queries.append(f"{product} {section_plan.title}")
        
        all_results = []
        seen_ids = set()
        
        for query in queries:
            results = self.hybrid_searcher.search(
                query=query,
                top_k=top_k,
                filter_dict=filter_dict
            )
            
            for r in results:
                if r.chunk_id not in seen_ids:
                    all_results.append(r)
                    seen_ids.add(r.chunk_id)
        
        # 重排序
        if self.reranker and all_results:
            combined_query = f"{section_plan.title} {' '.join(section_plan.required_topics)}"
            rerank_results = self.reranker.rerank(
                query=combined_query,
                results=all_results,
                top_k=top_k
            )
            
            contexts = [r.content for r in rerank_results]
            citations = [
                {
                    "chunk_id": r.chunk_id,
                    "doc_title": r.metadata.get("doc_title", ""),
                    "section": r.metadata.get("section_title", ""),
                    "page": r.metadata.get("page_start", 0),
                    "quote": r.content[:100]
                }
                for r in rerank_results
            ]
        else:
            contexts = [r.content for r in all_results[:top_k]]
            citations = [
                {
                    "chunk_id": r.chunk_id,
                    "doc_title": r.metadata.get("doc_title", ""),
                    "section": r.metadata.get("section_title", ""),
                    "page": r.metadata.get("page_start", 0),
                    "quote": r.content[:100]
                }
                for r in all_results[:top_k]
            ]
        
        return contexts, citations
    
    def _generate_section_content(
        self,
        section_plan: SectionPlan,
        contexts: List[str],
        citations: List[Dict],
        product: str,
        target_audience: str
    ) -> str:
        """生成章节内容"""
        if not self.llm_client:
            return f"[无法生成内容: LLM 客户端未配置]\n\n{section_plan.description}"
        
        # 构建上下文
        context_text = ""
        for i, (ctx, cite) in enumerate(zip(contexts, citations), 1):
            context_text += f"\n[资料{i}] 来源: {cite.get('doc_title', '')} - {cite.get('section', '')}\n{ctx}\n"
        
        if not context_text:
            context_text = "[无相关资料，请基于产品通用知识撰写]"
        
        prompt = f"""你是一位专业的技术文档撰写专家。请为以下章节撰写内容。

产品: {product}
目标读者: {target_audience}
章节标题: {section_plan.title}
章节描述: {section_plan.description}
预期字数: 约 {section_plan.word_count_estimate} 字

参考资料:
{context_text}

写作要求:
1. 内容必须基于提供的参考资料，不要编造信息
2. 使用清晰、专业的技术文档风格
3. 在使用参考资料中的信息时，标注来源 [资料X]
4. 如果是步骤说明，使用编号列表
5. 如果有参数或配置，使用表格格式
6. 适当使用代码块展示命令或代码示例

请直接输出章节内容 (不要包含标题，标题会自动添加):"""

        try:
            if hasattr(self, '_client_type') and self._client_type == "anthropic":
                response = self.llm_client.messages.create(
                    model=self.model,
                    max_tokens=2000,
                    messages=[{"role": "user", "content": prompt}]
                )
                return response.content[0].text.strip()
            else:
                response = self.llm_client.chat.completions.create(
                    model=self.model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,
                    max_tokens=2000
                )
                return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Failed to generate section content: {e}")
            return f"[生成失败: {str(e)}]\n\n{section_plan.description}"
